from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(13)
FONT_COLOR = RGBColor(0, 0, 0)

_test_steps = []
_test_results = {}
_test_cases = {}


# ============================================================
# COMMON HELPERS
# ============================================================

def get_current_time():
    return datetime.now().strftime("%d/%m/%Y %H:%M:%S")


def clean_xml_text(value):
    if value is None:
        return ""

    return re.sub(
        r"[\x00-\x08\x0B\x0C\x0E-\x1F]",
        "",
        str(value)
    )


def _set_run_format(run):
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.color.rgb = FONT_COLOR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def apply_report_format(document):
    """
    Áp dụng định dạng chung:
    Times New Roman - size 13 - màu đen.
    """

    normal_style = document.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = FONT_SIZE
    normal_style.font.color.rgb = FONT_COLOR
    normal_style._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        FONT_NAME
    )

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _set_run_format(run)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_format(run)


def _count_statuses(results):
    return {
        "total": len(results),
        "passed": sum(
            1 for result in results.values()
            if result["status"] == "PASSED"
        ),
        "failed": sum(
            1 for result in results.values()
            if result["status"] == "FAILED"
        ),
        "xfailed": sum(
            1 for result in results.values()
            if result["status"] == "XFAILED"
        ),
        "skipped": sum(
            1 for result in results.values()
            if result["status"] == "SKIPPED"
        )
    }


def _add_report_title(document, title_text):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run(title_text)
    title_run.bold = True

    generated_time = document.add_paragraph()
    generated_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated_time.add_run(
        f"Thời gian tạo báo cáo: {get_current_time()}"
    )

    document.add_paragraph()


def _add_summary_table(document, results):
    counts = _count_statuses(results)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    table.rows[0].cells[0].text = "Nội dung"
    table.rows[0].cells[1].text = "Kết quả"

    summary_data = [
        ("Tổng số Test Case", counts["total"]),
        ("Passed", counts["passed"]),
        ("Failed", counts["failed"]),
        ("XFailed", counts["xfailed"]),
        ("Skipped", counts["skipped"])
    ]

    for label, value in summary_data:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)


# ============================================================
# STEP REPORTING
# ============================================================

def report_step(
        test_case_id,
        step_number,
        description,
        status="PASS",
        detail=None
):
    recorded_at = get_current_time()

    message = (
        f"\n{test_case_id} | STEP {step_number} | "
        f"{status} | {description}"
    )

    if detail:
        message += f" | {detail}"

    print(message)

    _test_steps.append({
        "test_case_id": test_case_id,
        "step_number": step_number,
        "status": status,
        "description": description,
        "detail": detail or "",
        "time": recorded_at
    })


def report_failed_step(
        test_case_id,
        step_number,
        detail
):
    already_exists = any(
        step["test_case_id"] == test_case_id
        and str(step["step_number"]) == str(step_number)
        and step["status"] == "FAIL"
        for step in _test_steps
    )

    if already_exists:
        return

    report_step(
        test_case_id=test_case_id,
        step_number=step_number,
        description=f"Step {step_number} thực thi thất bại",
        status="FAIL",
        detail=detail
    )


# ============================================================
# TEST CASE RESULT
# ============================================================

def save_test_result(
        test_case_id,
        status,
        detail=None,
        duration=None,
        screenshot=None
):
    _test_results[test_case_id] = {
        "status": status,
        "detail": detail or "",
        "duration": duration,
        "screenshot": screenshot or "",
        "time": get_current_time()
    }


def register_test_case(
        test_case_id,
        feature_name,
        description=""
):
    clean_description = (description or "").strip()
    lines = clean_description.splitlines()

    if lines and lines[0].strip() == test_case_id:
        clean_description = "\n".join(lines[1:]).strip()

    _test_cases[test_case_id] = {
        "feature": feature_name,
        "description": clean_description
    }


def reset_test_report():
    _test_steps.clear()
    _test_results.clear()
    _test_cases.clear()


# ============================================================
# REPORT RIÊNG TỪNG CHỨC NĂNG
# ============================================================

def generate_word_report(
        output_path="reports/Notification_Test_Report.docx",
        feature_name="THÔNG BÁO",
        test_case_prefix=None
):
    output_path = Path(output_path)

    if not output_path.is_absolute():
        output_path = PROJECT_ROOT / output_path

    output_path.parent.mkdir(parents=True, exist_ok=True)

    if test_case_prefix:
        filtered_results = {
            test_case_id: result
            for test_case_id, result in _test_results.items()
            if test_case_id.startswith(test_case_prefix)
        }

        filtered_steps = [
            step
            for step in _test_steps
            if step["test_case_id"].startswith(test_case_prefix)
        ]
    else:
        filtered_results = _test_results
        filtered_steps = _test_steps

    print(
        f"\nREPORT DATA | Feature = {feature_name} | "
        f"Total TC = {len(filtered_results)} | "
        f"TC = {sorted(filtered_results.keys())}"
    )

    document = Document()

    _add_report_title(
        document,
        "BÁO CÁO KẾT QUẢ KIỂM THỬ TỰ ĐỘNG\n"
        f"CHỨC NĂNG {feature_name}"
    )

    # 1. Tổng quan
    document.add_heading("1. Tổng quan kết quả", level=1)
    _add_summary_table(document, filtered_results)
    document.add_paragraph()

    # 2. Kết quả từng TC
    document.add_heading("2. Kết quả từng Test Case", level=1)

    result_table = document.add_table(rows=1, cols=6)
    result_table.style = "Table Grid"

    headers = result_table.rows[0].cells
    headers[0].text = "Test Case"
    headers[1].text = "Kết quả"
    headers[2].text = "Thời gian"
    headers[3].text = "Duration (s)"
    headers[4].text = "Screenshot"
    headers[5].text = "Ghi chú"

    for test_case_id in sorted(filtered_results.keys()):
        result = filtered_results[test_case_id]
        row = result_table.add_row().cells

        row[0].text = test_case_id
        row[1].text = result["status"]
        row[2].text = result["time"]

        duration = result["duration"]
        row[3].text = (
            f"{duration:.2f}"
            if duration is not None
            else ""
        )

        row[4].text = clean_xml_text(result["screenshot"])
        row[5].text = clean_xml_text(result["detail"])

    document.add_paragraph()

    # 3. Chi tiết Step
    document.add_heading(
        "3. Chi tiết thực thi từng Step",
        level=1
    )

    step_table = document.add_table(rows=1, cols=6)
    step_table.style = "Table Grid"

    headers = step_table.rows[0].cells
    headers[0].text = "Test Case"
    headers[1].text = "Step"
    headers[2].text = "Kết quả"
    headers[3].text = "Thời gian"
    headers[4].text = "Mô tả"
    headers[5].text = "Chi tiết"

    for step in filtered_steps:
        row = step_table.add_row().cells

        row[0].text = step["test_case_id"]
        row[1].text = str(step["step_number"])
        row[2].text = step["status"]
        row[3].text = step["time"]
        row[4].text = clean_xml_text(step["description"])
        row[5].text = clean_xml_text(step["detail"])

    apply_report_format(document)
    document.save(output_path)

    print("\n" + "=" * 60)
    print(f"WORD TEST REPORT GENERATED | {feature_name}")
    print(output_path.resolve())
    print("=" * 60)

    return output_path


# ============================================================
# OVERALL TEST REPORT
# ============================================================

def generate_overall_report(
        output_path="reports/Overall_Test_Report.docx"
):
    output_path = Path(output_path)

    if not output_path.is_absolute():
        output_path = PROJECT_ROOT / output_path

    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()

    _add_report_title(
        document,
        "BÁO CÁO TỔNG HỢP KẾT QUẢ KIỂM THỬ TỰ ĐỘNG\n"
        "SELENIUM TEST SUITE"
    )

    # 1. Tổng quan
    document.add_heading("1. Tổng quan kết quả", level=1)
    _add_summary_table(document, _test_results)
    document.add_paragraph()

    # Gom TC theo chức năng
    feature_tests = {}

    for test_case_id, info in _test_cases.items():
        feature_name = info["feature"]

        feature_tests.setdefault(
            feature_name,
            []
        ).append(test_case_id)

    # 2. Tổng hợp theo chức năng
    document.add_heading(
        "2. Tổng hợp theo chức năng",
        level=1
    )

    feature_table = document.add_table(rows=1, cols=6)
    feature_table.style = "Table Grid"

    headers = feature_table.rows[0].cells
    headers[0].text = "Chức năng"
    headers[1].text = "Tổng TC"
    headers[2].text = "Passed"
    headers[3].text = "Failed"
    headers[4].text = "XFailed"
    headers[5].text = "Skipped"

    for feature_name in sorted(feature_tests.keys()):
        test_case_ids = feature_tests[feature_name]

        feature_results = {
            test_case_id: _test_results[test_case_id]
            for test_case_id in test_case_ids
            if test_case_id in _test_results
        }

        counts = _count_statuses(feature_results)
        row = feature_table.add_row().cells

        row[0].text = feature_name
        row[1].text = str(counts["total"])
        row[2].text = str(counts["passed"])
        row[3].text = str(counts["failed"])
        row[4].text = str(counts["xfailed"])
        row[5].text = str(counts["skipped"])

    document.add_paragraph()

    # 3. Chi tiết TC theo chức năng
    document.add_heading(
        "3. Chi tiết Test Case theo chức năng",
        level=1
    )

    for feature_name in sorted(feature_tests.keys()):
        test_case_ids = sorted(
            feature_tests[feature_name]
        )

        feature_results = {
            test_case_id: _test_results[test_case_id]
            for test_case_id in test_case_ids
            if test_case_id in _test_results
        }

        counts = _count_statuses(feature_results)

        document.add_heading(
            feature_name,
            level=2
        )

        document.add_paragraph(
            f"Tổng Test Case: {counts['total']} | "
            f"Passed: {counts['passed']} | "
            f"Failed: {counts['failed']} | "
            f"XFailed: {counts['xfailed']} | "
            f"Skipped: {counts['skipped']}"
        )

        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        headers = table.rows[0].cells
        headers[0].text = "Test Case"
        headers[1].text = "Mô tả"
        headers[2].text = "Kết quả"
        headers[3].text = "Ghi chú"

        for test_case_id in test_case_ids:
            if test_case_id not in _test_results:
                continue

            result = _test_results[test_case_id]
            info = _test_cases.get(test_case_id, {})

            row = table.add_row().cells

            row[0].text = test_case_id
            row[1].text = clean_xml_text(
                info.get("description", "")
            )
            row[2].text = result["status"]

            if result["status"] in (
                "FAILED",
                "XFAILED",
                "SKIPPED"
            ):
                row[3].text = clean_xml_text(
                    result["detail"]
                )
            else:
                row[3].text = ""

        document.add_paragraph()

    apply_report_format(document)
    document.save(output_path)

    print("\n" + "=" * 60)
    print("OVERALL TEST REPORT GENERATED:")
    print(output_path.resolve())
    print("=" * 60)

    return output_path


# ============================================================
# TEST CASE START
# ============================================================

def report_test_case_start(
        test_case_id,
        description
):
    print()
    print("=" * 70)
    print(test_case_id)
    print(description)
    print("=" * 70)